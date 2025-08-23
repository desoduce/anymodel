import io
import os
from typing import Dict, Any, Optional
import pandas as pd
import PyPDF2
from docx import Document
from fastapi import UploadFile
import csv
from prompt_cleaner import filterPII

class DocumentProcessor:
    """Process various document types and extract text content with PII filtering"""
    
    SUPPORTED_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
        'application/vnd.ms-excel': 'xls',
        'text/csv': 'csv',
        'text/plain': 'txt',
        'application/msword': 'doc'
    }
    
    async def process_file(self, file: UploadFile) -> Dict[str, Any]:
        """
        Process uploaded file and extract text content with PII filtering
        
        Args:
            file: FastAPI UploadFile object
            
        Returns:
            Dict containing processed content, metadata, and filtering info
        """
        try:
            # Read file content
            content = await file.read()
            await file.seek(0)  # Reset file pointer
            
            # Determine file type
            content_type = file.content_type
            file_extension = os.path.splitext(file.filename)[1].lower()
            
            # Extract text based on file type
            raw_text = ""
            metadata = {
                "filename": file.filename,
                "content_type": content_type,
                "size": len(content),
                "type": "unknown"
            }
            
            if content_type in self.SUPPORTED_TYPES:
                file_type = self.SUPPORTED_TYPES[content_type]
                metadata["type"] = file_type
                raw_text = await self._extract_text(content, file_type, file.filename)
            elif file_extension:
                # Try to determine by extension
                ext_map = {
                    '.pdf': 'pdf',
                    '.docx': 'docx',
                    '.xlsx': 'xlsx',
                    '.xls': 'xls',
                    '.csv': 'csv',
                    '.txt': 'txt'
                }
                if file_extension in ext_map:
                    file_type = ext_map[file_extension]
                    metadata["type"] = file_type
                    raw_text = await self._extract_text(content, file_type, file.filename)
            
            if not raw_text:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {content_type}",
                    "metadata": metadata
                }
            
            # Filter PII from extracted text
            filtered_text = filterPII(raw_text)
            
            # Calculate filtering statistics
            filtering_stats = self._calculate_filtering_stats(raw_text, filtered_text)
            
            return {
                "success": True,
                "raw_text": raw_text[:1000] + "..." if len(raw_text) > 1000 else raw_text,  # Truncated for security
                "filtered_text": filtered_text,
                "metadata": metadata,
                "filtering_stats": filtering_stats,
                "pii_detected": filtering_stats["items_filtered"] > 0
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "metadata": {"filename": file.filename if file else "unknown"}
            }
    
    async def _extract_text(self, content: bytes, file_type: str, filename: str) -> str:
        """Extract text from different file types"""
        
        if file_type == 'pdf':
            return self._extract_pdf_text(content)
        elif file_type == 'docx':
            return self._extract_docx_text(content)
        elif file_type in ['xlsx', 'xls']:
            return self._extract_excel_text(content)
        elif file_type == 'csv':
            return self._extract_csv_text(content)
        elif file_type == 'txt':
            return content.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error reading PDF: {str(e)}")
    
    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {str(e)}")
    
    def _extract_excel_text(self, content: bytes) -> str:
        """Extract text from Excel files"""
        try:
            excel_file = io.BytesIO(content)
            
            # Read all sheets
            excel_data = pd.read_excel(excel_file, sheet_name=None)
            
            text = ""
            for sheet_name, df in excel_data.items():
                text += f"Sheet: {sheet_name}\n"
                text += df.to_string(index=False) + "\n\n"
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error reading Excel: {str(e)}")
    
    def _extract_csv_text(self, content: bytes) -> str:
        """Extract text from CSV"""
        try:
            csv_content = content.decode('utf-8', errors='ignore')
            lines = csv_content.split('\n')
            
            # Read CSV data
            csv_reader = csv.reader(lines)
            text = ""
            
            for row in csv_reader:
                text += ", ".join(row) + "\n"
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error reading CSV: {str(e)}")
    
    def _calculate_filtering_stats(self, original: str, filtered: str) -> Dict[str, Any]:
        """Calculate statistics about PII filtering"""
        
        # Count filtered items by looking for filter markers
        filter_markers = [
            '[SSN_FILTERED]', '[PHONE_FILTERED]', '[EMAIL_FILTERED]',
            '[CARD_FILTERED]', '[ADDRESS_FILTERED]', '[ZIP_FILTERED]',
            '[IP_FILTERED]', '[ACCOUNT_FILTERED]', '[DL_FILTERED]',
            '[PASSPORT_FILTERED]'
        ]
        
        items_filtered = 0
        filter_types = []
        
        for marker in filter_markers:
            count = filtered.count(marker)
            if count > 0:
                items_filtered += count
                filter_type = marker.replace('[', '').replace('_FILTERED]', '').lower()
                filter_types.append(f"{filter_type} ({count})")
        
        # Count name replacements (look for initials pattern)
        # Count instances where full names were likely replaced with initials
        import re
        initial_pattern = r'\b[A-Z]\.[A-Z]\.'
        name_replacements = len(re.findall(initial_pattern, filtered))
        
        # Estimate name filtering by comparing word counts of capitalized sequences
        original_caps = len(re.findall(r'\b[A-Z][a-z]{2,}\s+[A-Z][a-z]{2,}\b', original))
        filtered_caps = len(re.findall(r'\b[A-Z][a-z]{2,}\s+[A-Z][a-z]{2,}\b', filtered))
        likely_names_filtered = max(0, original_caps - filtered_caps)
        
        if likely_names_filtered > 0 or name_replacements > 0:
            items_filtered += max(likely_names_filtered, name_replacements)
            filter_types.append(f"names ({max(likely_names_filtered, name_replacements)})")
        
        return {
            "original_length": len(original),
            "filtered_length": len(filtered),
            "reduction_percentage": round((len(original) - len(filtered)) / len(original) * 100, 2) if original else 0,
            "items_filtered": items_filtered,
            "filter_types": filter_types
        }

    def get_supported_extensions(self) -> list:
        """Get list of supported file extensions"""
        return ['.pdf', '.docx', '.xlsx', '.xls', '.csv', '.txt']
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if file type is supported"""
        ext = os.path.splitext(filename)[1].lower()
        return ext in self.get_supported_extensions()
